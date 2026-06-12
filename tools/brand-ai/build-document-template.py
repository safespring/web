from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "static" / "brand" / "ai" / "templates" / "safespring-document-template.docx"
LOGO_BLUE = ROOT / "static" / "img" / "logos" / "safespring" / "png" / "safespring_logotype_blue_png.png"

COLORS = {
    "blue": "195F8C",
    "middle_blue": "417DA5",
    "clear_blue": "3C9BCD",
    "green": "32CD32",
    "web_green": "19F064",
    "orange": "FA690F",
    "cloud_blue": "E7EFF3",
    "code_blue": "EAF0F4",
    "charcoal": "323232",
    "cloud_white": "FAFEFE",
    "white": "FFFFFF",
    "component_bg": "E7EFF3",
}

PAIRS = {
    "blue": {"bg": COLORS["component_bg"], "fg": COLORS["blue"]},
    "middle_blue": {"bg": COLORS["cloud_blue"], "fg": COLORS["middle_blue"]},
    "clear_blue": {"bg": COLORS["cloud_blue"], "fg": COLORS["clear_blue"]},
    "green": {"bg": COLORS["cloud_blue"], "fg": COLORS["green"]},
    "web_green": {"bg": COLORS["cloud_blue"], "fg": COLORS["web_green"]},
    "orange": {"bg": COLORS["cloud_blue"], "fg": COLORS["orange"]},
    "neutral": {"bg": COLORS["cloud_blue"], "fg": COLORS["charcoal"]},
}


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="C9DCE6", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = f"w:{name}"
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color, size="18"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "10")
    bottom.set(qn("w:color"), color)


def set_style_font(style, family, size=None, color=None, bold=None):
    style.font.name = family
    style._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold


def add_text(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    return run


def configure_styles(doc):
    styles = doc.styles
    set_style_font(styles["Normal"], "Montserrat", 10.5, COLORS["charcoal"])
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Title"], "Hind", 34, COLORS["blue"], False)
    styles["Title"].paragraph_format.space_after = Pt(12)

    set_style_font(styles["Subtitle"], "Montserrat", 13, "52626A", False)
    styles["Subtitle"].paragraph_format.space_after = Pt(18)

    set_style_font(styles["Heading 1"], "Hind", 24, COLORS["blue"], False)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Heading 2"], "Montserrat", 14, COLORS["clear_blue"], True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    set_style_font(styles["Heading 3"], "Montserrat", 11, COLORS["blue"], True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.add_run().add_picture(str(LOGO_BLUE), width=Inches(1.55))

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Safespring AI Brand Kit | safespring.com/brand/ai")
    footer_run.font.name = "Safespring Mono"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = rgb("6A7B83")


def add_callout(doc, label, text, pair_name="blue"):
    pair = PAIRS[pair_name]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, pair["bg"])
    set_cell_borders(cell, "C9DCE6")
    set_cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_text(p, label.upper(), True, pair["fg"])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    add_text(p2, text, False, pair["fg"])
    doc.add_paragraph()


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, COLORS["code_blue"])
    set_cell_borders(cell, "C9DCE6")
    set_cell_margins(cell, 150, 190, 150, 190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    code_run = p.add_run(text)
    code_run.font.name = "Safespring Mono"
    code_run.font.size = Pt(9)
    code_run.font.color.rgb = rgb(COLORS["blue"])


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    add_header_footer(section)

    title = doc.add_paragraph(style="Title")
    title.add_run("Document title")
    set_paragraph_border(title, COLORS["orange"])

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Subtitle or decision statement. Use this area to state the document purpose clearly.")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(18)
    meta.paragraph_format.space_after = Pt(40)
    meta_run = meta.add_run("Prepared for [audience] | Version 1.0 | YYYY-MM-DD")
    meta_run.font.name = "Safespring Mono"
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = rgb("52626A")

    add_callout(
        doc,
        "Cloud-blue background rule",
        "Use cloud blue as the background for text boxes, callouts, tables, and graphics. Main colors are for text, icons, lines, and emphasis.",
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "Open with the main claim, decision, or recommendation. Keep paragraphs short and use evidence only where it helps the reader understand what to do next."
    )
    doc.add_paragraph("Use this template as a starting point for AI-generated or manually written Safespring documents.")

    doc.add_heading("Recommended structure", level=2)
    for item in [
        "Context: what changed and why the reader should care.",
        "Evidence: architecture, requirements, operational controls, or source-backed facts.",
        "Action: owner, next step, timeline, and decision point.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Example table", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.autofit = False
    table.allow_autofit = False
    widths = [Inches(1.55), Inches(2.45), Inches(2.45)]
    headers = ["Area", "Use", "Avoid"]
    rows = [
        ["middle_blue", "Security", "Specific controls and responsibilities.", "Fear-based or generic security imagery."],
        ["green", "Compliance", "Precise source-backed statements.", "Invented certifications or claims."],
        ["orange", "Cloud", "Concrete architecture and boundaries.", "Abstract neon cloud metaphors."],
    ]
    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.width = widths[col]
        set_cell_shading(cell, PAIRS["blue"]["bg"])
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header, True, PAIRS["blue"]["fg"])
    for row_index, row in enumerate(rows, start=1):
        pair = PAIRS[row[0]]
        for col, value in enumerate(row[1:]):
            cell = table.cell(row_index, col)
            cell.width = widths[col]
            set_cell_shading(cell, pair["bg"])
            set_cell_borders(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, col == 0, pair["fg"])

    doc.add_paragraph()
    doc.add_heading("Technical block", level=2)
    add_code_block(doc, "brand-kit: https://www.safespring.com/brand/ai/brand-kit.json")

    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("Use appendices for sources, detailed requirements, or long technical tables.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
